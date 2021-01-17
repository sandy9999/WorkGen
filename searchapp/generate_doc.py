from docx import Document
import datetime
import os
from django.conf import settings
from json import dumps, loads
import pickle
from .models import FormAuth
import os.path
from googleapiclient import errors
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request

qtype_to_section = {"1A": "A", "1B": "B", "2": "C", "3": "D", "5": "E"}


SCOPES = ['https://www.googleapis.com/auth/forms']
API_ID="MisrDt1oXa2FhUFiDHl9Yn5oDBuBhNfqu"

def get_no_of_questions(questions_mapping, question_type):
    count = 0
    for row in questions_mapping:
        if row.get('question_type') == question_type:
            for question in row['question']:
                count = count + 1
    return count


def add_headings(document, subject):
    heading = document.add_heading("EXAM")
    heading.alignment = 1
    subject_heading = document.add_heading(subject, level=2)
    subject_heading.alignment = 1
    para = document.add_paragraph()
    run = para.add_run()
    run.add_break()


def add_questions(document, questions_mapping, question_type, question_no, no_of_questions):
    section = document.add_paragraph()
    section.add_run("Section " + qtype_to_section[question_type]).bold = True
    section.alignment = 1
    subsection = document.add_paragraph()
    subsection.add_run(str(no_of_questions) + " questions of " + question_type[0] + " marks each")
    subsection.alignment = 1
    for row in questions_mapping:
        if row.get('question_type') == question_type:
            p = document.add_paragraph()
            p.add_run("Attempt only " + str(min(row['attempt'], no_of_questions)) + " questions").italic = True
            p.alignment = 1
            for question in row['question']:
                question_no = question_no + 1
                p1 = document.add_paragraph(str(question_no) + ": ")
                p1.add_run(question)
    para = document.add_paragraph()
    run = para.add_run()
    run.add_break()
    return question_no


def convert_customized_to_doc(questions_mapping, subject):
    document = Document()
    for name in questions_mapping:
        para = document.add_paragraph()
        student_name = para.add_run("Name: " + name)
        student_name.add_break()
        add_headings(document, subject)
        question_no = 0
        for qtype in qtype_to_section:
            no_of_questions = get_no_of_questions(questions_mapping[name], qtype)
            if no_of_questions != 0:
                question_no = add_questions(document, questions_mapping[name], qtype, question_no, no_of_questions)
        document.add_page_break()
    filename = "question_worksheet_" + datetime.datetime.now().__str__() + ".docx"
    filepath = os.path.join('searchapp/static/docs', filename)
    document.save(filepath)
    return filepath


def convert_to_doc(questions_mapping, subject):
    document = Document()
    add_headings(document, subject)
    question_no = 0
    for qtype in qtype_to_section:
        no_of_questions = get_no_of_questions(questions_mapping, qtype)
        if no_of_questions != 0:
            question_no = add_questions(document, questions_mapping, qtype, question_no, no_of_questions)
    filename = "question_worksheet_" + datetime.datetime.now().__str__() + ".docx"
    filepath = os.path.join('searchapp/static/docs', filename)
    document.save(filepath)
    return filepath

def convert_to_form(questions_mapping, subject, required, heading, user):
    creds = get_creds(user)
    service = build('script', 'v1', credentials=creds) 
    required = dumps(required)
    questions_mapping = dumps(questions_mapping)
    request = {"function": "create_form", "parameters": [{"required": required, "heading": heading, "question_mapping": questions_mapping, "subject": subject}], "devMode": "true"}
    response = service.scripts().run(body=request, scriptId=API_ID).execute()
    return response

def get_creds(user):
    creds = None
    flag = 0
    if FormAuth.objects.filter(user=user).exists():
        data = FormAuth.objects.get(user=user).creds
        creds = pickle.loads(data)
        flag = 1
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                    'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
            if(flag == 1):
                token = FormAuth.objects.get(user=user)
                token.creds = pickle.dumps(creds)
                token.save()
            else:
                data = FormAuth(user=user, creds=pickle.dumps(creds))
                data.save()
    return creds
